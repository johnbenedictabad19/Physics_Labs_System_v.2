from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import base64

# ============================================================
# MATH CONVERSION TABLES  (OMML → LaTeX)
# ============================================================

# Unicode characters → LaTeX commands
_UNICODE_TO_LATEX = {
    # Invisible / zero-width chars — remove these
    '\u2061': '', '\u2062': '', '\u2063': '', '\u2064': '',
    '\u200b': '', '\u200c': '', '\u200d': '', '\ufeff': '',
    # Common math symbols
    '∞': r'\infty',   '±': r'\pm',       '∓': r'\mp',
    '×': r'\times',   '÷': r'\div',      '·': r'\cdot',
    '∑': r'\sum',     '∏': r'\prod',     '∫': r'\int',
    '∂': r'\partial', '∇': r'\nabla',    '∆': r'\Delta',
    '√': r'\sqrt',
    '≤': r'\leq',     '≥': r'\geq',      '≠': r'\neq',
    '≈': r'\approx',  '≡': r'\equiv',    '∝': r'\propto',
    '→': r'\rightarrow',  '←': r'\leftarrow',  '↔': r'\leftrightarrow',
    '⇒': r'\Rightarrow',  '⇔': r'\Leftrightarrow',
    '∈': r'\in',  '∉': r'\notin',  '⊂': r'\subset',  '⊃': r'\supset',
    '∪': r'\cup',  '∩': r'\cap',  '∅': r'\emptyset',
    '°': r'^\circ',
    # Greek lowercase
    'α': r'\alpha',  'β': r'\beta',   'γ': r'\gamma',  'δ': r'\delta',
    'ε': r'\epsilon','ζ': r'\zeta',   'η': r'\eta',    'θ': r'\theta',
    'ι': r'\iota',   'κ': r'\kappa',  'λ': r'\lambda', 'μ': r'\mu',
    'ν': r'\nu',     'ξ': r'\xi',     'π': r'\pi',     'ρ': r'\rho',
    'σ': r'\sigma',  'τ': r'\tau',    'υ': r'\upsilon','φ': r'\phi',
    'χ': r'\chi',    'ψ': r'\psi',    'ω': r'\omega',
    # Greek uppercase
    'Γ': r'\Gamma',  'Δ': r'\Delta',  'Θ': r'\Theta',  'Λ': r'\Lambda',
    'Ξ': r'\Xi',     'Π': r'\Pi',     'Σ': r'\Sigma',  'Υ': r'\Upsilon',
    'Φ': r'\Phi',    'Ψ': r'\Psi',    'Ω': r'\Omega',
}

# Common function names → LaTeX commands
_MATH_FUNCS = {
    'sin': r'\sin',    'cos': r'\cos',    'tan': r'\tan',
    'csc': r'\csc',    'sec': r'\sec',    'cot': r'\cot',
    'arcsin': r'\arcsin', 'arccos': r'\arccos', 'arctan': r'\arctan',
    'sinh': r'\sinh',  'cosh': r'\cosh',  'tanh': r'\tanh',
    'ln':  r'\ln',     'log': r'\log',    'exp': r'\exp',
    'lim': r'\lim',    'max': r'\max',    'min': r'\min',
    'det': r'\det',    'dim': r'\dim',    'ker': r'\ker',
    'gcd': r'\gcd',    'mod': r'\bmod',
}

# ============================================================
# HELPERS
# ============================================================
# Build the symbol-char set now that _UNICODE_TO_LATEX is complete
_MATH_SYMBOL_CHARS = frozenset(k for k, v in _UNICODE_TO_LATEX.items() if v)

def extract_images(doc):
    """Extract all images from doc relationships -> {rId: {data, ext}}

    Scans both the main document part and any embedded chart/drawing parts
    so images inside shapes/frames are also captured.
    """
    images = {}

    def _add_from_part(part):
        for rel in part.rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                img_part = rel.target_part
                img_data = img_part.blob
                img_ext  = img_part.content_type.split('/')[-1]
                if img_ext == 'jpeg':
                    img_ext = 'jpg'
                b64 = base64.b64encode(img_data).decode('utf-8')
                images[rel.rId] = {
                    'data': f"data:image/{img_ext};base64,{b64}",
                    'ext':  img_ext
                }
            except Exception:
                pass

    _add_from_part(doc.part)
    # Also scan drawing parts (charts, diagrams, etc.)
    for rel in doc.part.rels.values():
        try:
            if hasattr(rel, 'target_part') and hasattr(rel.target_part, 'rels'):
                _add_from_part(rel.target_part)
        except Exception:
            pass

    return images


def extract_math_latex(omath_el):
    """Convert OMML m:oMath element to a LaTeX string for MathJax rendering."""

    # Delimiter char → (LaTeX \left cmd, LaTeX \right cmd)
    _DELIM = {
        '(':  (r'\left(',    r'\right)'),
        '[':  (r'\left[',   r'\right]'),
        '{':  (r'\left\{',  r'\right\}'),
        '|':  (r'\left|',   r'\right|'),
        '‖':  (r'\left\|',  r'\right\|'),
        '⌊':  (r'\left\lfloor', r'\right\rfloor'),
        '⌈':  (r'\left\lceil',  r'\right\rceil'),
    }

    def escape_t(text):
        result = ''
        for ch in text:
            mapped = _UNICODE_TO_LATEX.get(ch)
            if mapped is not None:
                result += mapped  # '' strips invisible chars; '\alpha' etc. for symbols
            else:
                result += ch
        return result

    def proc(el):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag

        if tag == 't':
            return escape_t(el.text or '')

        elif tag in ('r', 'oMath', 'oMathPara', 'box',
                     'e', 'fName', 'limLow', 'limUpp', 'lim',
                     'groupChr', 'bar', 'acc'):
            return proc_children(el)

        elif tag == 'f':  # fraction → \frac{num}{den}
            num = proc_children(el.find(qn('m:num')))
            den = proc_children(el.find(qn('m:den')))
            return r'\frac{' + num + '}{' + den + '}'

        elif tag == 'sSup':  # superscript → base^{exp}
            base = proc_children(el.find(qn('m:e')))
            sup  = proc_children(el.find(qn('m:sup')))
            return '{' + base + '}^{' + sup + '}'

        elif tag == 'sSub':  # subscript → base_{sub}
            base = proc_children(el.find(qn('m:e')))
            sub  = proc_children(el.find(qn('m:sub')))
            return '{' + base + '}_{' + sub + '}'

        elif tag == 'sSubSup':  # subscript + superscript
            base = proc_children(el.find(qn('m:e')))
            sub  = proc_children(el.find(qn('m:sub')))
            sup  = proc_children(el.find(qn('m:sup')))
            return '{' + base + '}_{' + sub + '}^{' + sup + '}'

        elif tag == 'rad':  # radical → \sqrt[n]{base}
            deg  = proc_children(el.find(qn('m:deg')))
            base = proc_children(el.find(qn('m:e')))
            if deg.strip():
                return r'\sqrt[' + deg + ']{' + base + '}'
            return r'\sqrt{' + base + '}'

        elif tag == 'nary':  # n-ary operator (∑ ∫ ∏ …)
            naryPr  = el.find(qn('m:naryPr'))
            char    = r'\int'
            if naryPr is not None:
                chr_el = naryPr.find(qn('m:chr'))
                if chr_el is not None:
                    raw = chr_el.get(qn('m:val'), '∫')
                    char = _UNICODE_TO_LATEX.get(raw, raw)
            sub  = proc_children(el.find(qn('m:sub')))
            sup  = proc_children(el.find(qn('m:sup')))
            body = proc_children(el.find(qn('m:e')))
            result = char
            if sub:
                result += '_{' + sub + '}'
            if sup:
                result += '^{' + sup + '}'
            return result + ' ' + body

        elif tag == 'func':  # function application → \sin arg
            fname_raw = proc_children(el.find(qn('m:fName'))).strip()
            arg       = proc_children(el.find(qn('m:e')))
            fname     = _MATH_FUNCS.get(fname_raw, fname_raw)
            return fname + ' ' + arg

        elif tag == 'd':  # delimiter → \left( inner \right)
            dPr = el.find(qn('m:dPr'))
            beg_cmd, end_cmd = r'\left(', r'\right)'
            if dPr is not None:
                b_el = dPr.find(qn('m:begChr'))
                e_el = dPr.find(qn('m:endChr'))
                bch  = b_el.get(qn('m:val'), '(') if b_el is not None else '('
                ech  = e_el.get(qn('m:val'), ')') if e_el is not None else ')'
                pair = _DELIM.get(bch)
                if pair:
                    beg_cmd, end_cmd = pair
                else:
                    beg_cmd = r'\left'  + bch
                    end_cmd = r'\right' + ech
            parts = [proc_children(e) for e in el.findall(qn('m:e'))]
            return beg_cmd + ', '.join(parts) + end_cmd

        elif tag == 'eqArr':  # equation array — rows joined with \\
            rows = [proc_children(r) for r in el.findall(qn('m:e'))]
            return r' \\ '.join(rows)

        elif tag.endswith('Pr'):  # property elements — skip
            return ''

        else:
            return proc_children(el)

    def proc_children(el):
        if el is None:
            return ''
        return ''.join(proc(child) for child in el)

    return proc_children(omath_el).strip()


_VML_NS  = 'urn:schemas-microsoft-com:vml'
_REL_NS  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

def get_para_image_ids(para):
    """Return list of rIds of images embedded in a paragraph.

    Handles both DrawingML (a:blip r:embed) and legacy VML (v:imagedata r:id).
    """
    ids = []
    # DrawingML images
    for blip in para._element.findall('.//' + qn('a:blip')):
        embed = blip.get(qn('r:embed'))
        if embed:
            ids.append(embed)
    # VML images (older Word format)
    vml_tag = f'{{{_VML_NS}}}imagedata'
    rel_id_attr = f'{{{_REL_NS}}}id'
    for imgdata in para._element.findall('.//' + vml_tag):
        rid = imgdata.get(rel_id_attr)
        if rid:
            ids.append(rid)
    return ids


_INVISIBLE_MATH = str.maketrans('', '', '\u2061\u2062\u2063\u2064\u200b\u200c\u200d\ufeff')

# Chars that indicate a token is a Unicode Math expression
_UNICODE_MATH_CHARS = frozenset(
    list('\u200b')           # ZWS subscript delimiter
    + [k for k, v in _UNICODE_TO_LATEX.items() if v and ord(k) > 127]  # Greek + math symbols
)


def _apply_inline_math(text: str) -> str:
    """
    Post-process a text string: find contiguous non-space tokens that contain
    Unicode Math chars (ZWS, Greek letters, math symbols) and convert them to
    LaTeX inline \\(...\\).  Tokens already inside \\(...\\) or \\[...\\] are left alone.
    """
    # Don't touch text that's already wrapped in LaTeX delimiters
    if text.startswith('\\[') or text.startswith('\\('):
        return text

    def convert_token(m):
        token = m.group(0)
        if any(ch in _UNICODE_MATH_CHARS for ch in token):
            return f'\\({unicode_math_to_latex(token)}\\)'
        return token.translate(_INVISIBLE_MATH)

    # Match each non-space sequence; skip regions already inside \(...\) or \[...\]
    # Simple approach: process token by token
    parts = re.split(r'(\\\(.*?\\\)|\\\[.*?\\\])', text, flags=re.DOTALL)
    result = []
    for part in parts:
        if part.startswith('\\(') or part.startswith('\\['):
            result.append(part)   # already LaTeX — leave it
        else:
            result.append(re.sub(r'[^ ]+', convert_token, part))
    return ''.join(result)


def unicode_math_to_latex(text: str) -> str:
    """
    Convert Unicode Math linear form extracted from OMML to proper LaTeX.

    Handles three transformations in order:
      1. ZWS-delimited subscripts  n1U+200B  →  n_{1}
      2. Unicode Greek/math chars  θ → \\theta, ≤ → \\leq, etc.
      3. Bare function names       sin → \\sin, cos → \\cos, etc.
    """
    # 1. ZWS-bounded subscripts: [letter or Greek][digits]U+200B → \1_{\2}
    text = re.sub(
        r'([a-zA-Z\u0391-\u03c9\u03d0-\u03d6])(\d+)\u200b',
        r'\1_{\2}',
        text
    )
    # 2. Strip remaining invisible/zero-width chars
    text = text.translate(_INVISIBLE_MATH)
    # 3. Map Unicode symbols → LaTeX commands (char by char)
    out = []
    for ch in text:
        mapped = _UNICODE_TO_LATEX.get(ch)
        if mapped is not None:
            out.append(mapped)   # '' strips invisible; '\alpha' etc. for symbols
        else:
            out.append(ch)
    text = ''.join(out)
    # 4. Convert bare function names → LaTeX commands (longest match first)
    # Use lambda replacement to avoid re.sub misinterpreting \s, \t, etc.
    for fname, cmd in sorted(_MATH_FUNCS.items(), key=lambda x: -len(x[0])):
        text = re.sub(
            r'(?<!\\)(?<![a-zA-Z])' + re.escape(fname) + r'(?![a-zA-Z])',
            lambda m, _cmd=cmd: _cmd, text
        )
    return text.strip()


def _is_pure_equation(text: str) -> bool:
    """
    Return True when *text* looks like a standalone Unicode Math equation.

    Criteria:
      • Contains at least one Unicode math/Greek character (not in ASCII range).
      • After stripping invisible chars, the text has NO regular spaces
        (pure equations like n₁sinθ₁=n₂sinθ₂ have no prose spaces).
    """
    stripped = text.translate(_INVISIBLE_MATH)
    if not stripped:
        return False
    has_math_char = any(ch in _MATH_SYMBOL_CHARS for ch in stripped)
    has_space = ' ' in stripped
    return has_math_char and not has_space

def _run_text(wr_el):
    """Extract clean text from a w:r element, stripping invisible math chars."""
    t_el = wr_el.find(qn('w:t'))
    if t_el is None or not t_el.text:
        return ''
    return t_el.text.translate(_INVISIBLE_MATH)


def _run_script(wr_el):
    """Return 'sub', 'sup', or None based on w:rPr/w:vertAlign."""
    rPr = wr_el.find(qn('w:rPr'))
    if rPr is None:
        return None
    va = rPr.find(qn('w:vertAlign'))
    if va is None:
        return None
    val = va.get(qn('w:val'), '')
    if val == 'subscript':
        return 'sub'
    if val == 'superscript':
        return 'sup'
    return None


def get_full_para_text(para):
    """
    Extract paragraph text with equations as LaTeX for MathJax.

    Block equations (m:oMathPara):
      Word stores 3 alt-text forms alongside the OMML.  Form 2 is LaTeX
      (contains backslash commands like \\sin, \\theta).  We use Form 2 when
      available; otherwise fall back to structured OMML extraction.

    Inline equations (m:oMath):
      Processed by extract_math_latex → wrapped in \\(...\\).

    Inline subscripts/superscripts in prose (w:vertAlign):
      w:r runs with subscript/superscript character formatting are wrapped
      in \\(...\\) so MathJax renders them correctly.
    """
    el = para._element

    # ── BLOCK EQUATION PARAGRAPH ────────────────────────────────────────────
    # Use findall('.//')  so we catch oMathPara even inside w:sdt wrappers.
    has_block_eq = bool(el.find('.//' + qn('m:oMathPara')))

    if has_block_eq:
        # Try every w:r descendant for a LaTeX-form alt-text (has backslash cmds)
        latex_parts = []
        for wr in el.findall('.//' + qn('w:r')):
            txt = _run_text(wr)
            if txt and re.search(r'\\[a-zA-Z]', txt):
                latex_parts.append(txt)

        if latex_parts:
            latex = ''.join(latex_parts).strip()
            return f'\\[{latex}\\]'

        # Fallback A: collect raw w:r text (preserve ZWS for subscript detection)
        # → unicode_math_to_latex handles the ZWS → _{n} conversion
        wr_parts_raw = []
        for wr in el.findall('.//' + qn('w:r')):
            t_el = wr.find(qn('w:t'))
            if t_el is not None and t_el.text:
                wr_parts_raw.append(t_el.text)
        wr_text_raw = ''.join(wr_parts_raw)
        if _is_pure_equation(wr_text_raw):
            return f'\\[{unicode_math_to_latex(wr_text_raw)}\\]'

        # Fallback B: structured OMML extraction
        parts = []
        for omath in el.findall('.//' + qn('m:oMath')):
            eq = extract_math_latex(omath)
            if eq:
                eq = unicode_math_to_latex(eq)   # clean up any residual Unicode math
                parts.append(f'\\[{eq}\\]')
        return ''.join(parts).strip()

    # ── REGULAR / INLINE-EQUATION PARAGRAPH ─────────────────────────────────
    parts = []
    # We track pending script runs so consecutive sub/superscripts are grouped.
    pending_script = []   # list of (kind, text) where kind = 'sub'|'sup'

    def flush_script():
        if not pending_script:
            return
        # Build a minimal inline LaTeX expression: base_{sub}^{sup}
        # Group consecutive same-kind runs; mix is unusual but handle gracefully.
        expr = ''
        for kind, txt in pending_script:
            if kind == 'sub':
                expr += f'_{{{txt}}}'
            else:
                expr += f'^{{{txt}}}'
        parts.append(f'\\({expr}\\)')
        pending_script.clear()

    for child in el:   # direct children of w:p
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'r':
            t_el = child.find(qn('w:t'))
            raw_txt = t_el.text if t_el is not None and t_el.text else ''
            kind = _run_script(child)
            if not raw_txt.translate(_INVISIBLE_MATH):   # skip fully-invisible runs
                continue
            if kind:
                pending_script.append((kind, raw_txt.translate(_INVISIBLE_MATH)))
            else:
                flush_script()
                parts.append(raw_txt)   # keep ZWS intact for _apply_inline_math

        elif tag == 'hyperlink':
            flush_script()
            for wr in child.findall(qn('w:r')):
                txt = _run_text(wr)
                if txt:
                    parts.append(txt)

        elif tag == 'oMath':   # inline equation
            flush_script()
            eq = extract_math_latex(child)
            if eq:
                eq = unicode_math_to_latex(eq)
                parts.append(f'\\({eq}\\)')

        else:
            flush_script()
            # Catch m:oMath nested inside w:sdt or any other wrapper element
            for nested in child.findall('.//' + qn('m:oMath')):
                eq = extract_math_latex(nested)
                if eq:
                    eq = unicode_math_to_latex(eq)
                    parts.append(f'\\({eq}\\)')

    flush_script()
    result = ''.join(parts).strip()

    # If the whole paragraph is a pure Unicode math expression, wrap as block.
    if result and not result.startswith('\\') and _is_pure_equation(result):
        return f'\\[{unicode_math_to_latex(result)}\\]'

    # Final pass: convert any remaining Unicode Math tokens embedded in prose
    # (equations stored as w:r plain text rather than OMML).
    result = _apply_inline_math(result)

    return result


def is_bold(paragraph):
    """Return True if all non-empty runs are bold (explicit or style-inherited)."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    # r.bold is True (explicit), False (explicit off), or None (inherit from style)
    if all(r.bold is True for r in runs):
        return True
    if any(r.bold is False for r in runs):
        return False
    # Some/all runs inherit — check paragraph style
    try:
        if paragraph.style and paragraph.style.font.bold:
            return True
    except Exception:
        pass
    return False


_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT:    'left',
    WD_ALIGN_PARAGRAPH.CENTER:  'center',
    WD_ALIGN_PARAGRAPH.RIGHT:   'right',
    WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
}

def _para_style(para):
    """Return dict with 'align' and 'indent' keys for a paragraph."""
    align = _ALIGN_MAP.get(para.alignment, 'justify')  # default justify
    fmt = para.paragraph_format
    first_indent = fmt.first_line_indent
    # Positive first_line_indent = indented first line (like a tab)
    indent = first_indent is not None and first_indent > 0
    return {'align': align, 'indent': indent}


def _has_underline_run(para):
    """Return True if any run in the paragraph has underline formatting."""
    for run in para.runs:
        rPr = run._element.find('.//' + qn('w:rPr'))
        if rPr is not None:
            u = rPr.find(qn('w:u'))
            if u is not None and u.get(qn('w:val'), 'none') not in ('none', ''):
                return True
    return False


def _parse_ds_field_para(para):
    """
    Detect label+blank structure in a data_sheet paragraph.
    Blanks are either:
      - 3+ consecutive underscore characters (___)
      - Runs with underline formatting and whitespace-only content (Word-style fill-in-the-blank)
    Returns list of {'kind': 'label'/'blank', 'text': str} parts,
    or None if no blank detected.
    """
    el = para._element
    parts = []
    current_label = []
    has_blank = False

    def _flush_label():
        text = ''.join(current_label).strip()
        current_label.clear()
        if text:
            parts.append({'kind': 'label', 'text': text})

    for wr in el.findall('.//' + qn('w:r')):
        t_el = wr.find(qn('w:t'))
        raw_txt = t_el.text if t_el is not None and t_el.text else ''
        if not raw_txt:
            continue

        # Check for literal underscores (3+)
        if re.search(r'_{3,}', raw_txt):
            segments = re.split(r'_{3,}', raw_txt)
            for i, seg in enumerate(segments):
                if seg.strip():
                    current_label.append(seg)
                if i < len(segments) - 1:
                    _flush_label()
                    parts.append({'kind': 'blank'})
                    has_blank = True
            continue

        # Check for underline-formatted run (Word fill-in-the-blank)
        rpr = wr.find(qn('w:rPr'))
        is_underlined = False
        if rpr is not None:
            u_el = rpr.find(qn('w:u'))
            if u_el is not None:
                val = u_el.get(qn('w:val'), 'none')
                is_underlined = val not in ('none', 'noStrike', '')

        if is_underlined and not raw_txt.strip():
            # Underlined whitespace = blank
            _flush_label()
            if not parts or parts[-1]['kind'] != 'blank':
                parts.append({'kind': 'blank'})
                has_blank = True
        else:
            current_label.append(raw_txt)

    _flush_label()

    if not has_blank or not parts:
        return None
    return parts


# Matches letter-prefix group headers in procedures: "A.", "B.", "C.", "A. Title"
_PROC_LETTER_HDR_RE = re.compile(r'^[A-Z]\.\s', re.IGNORECASE)

# Matches numbered or lettered headers in data_sheet: "1. Title", "A. Title", "B. Results"
_DS_HDR_PREFIX_RE = re.compile(r'^([A-Z]|\d+)\.\s+\S', re.IGNORECASE)

# numFmt values that indicate a letter/roman list (→ group header, not a step)
_HEADER_NUM_FMTS = {'upperLetter', 'lowerLetter', 'upperRoman', 'lowerRoman'}

def _int_to_roman(n):
    vals = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
    syms = ['M','CM','D','CD','C','XC','L','XL','X','IX','V','IV','I']
    result = ''
    for v, s in zip(vals, syms):
        while n >= v:
            result += s
            n -= v
    return result

def _make_list_label(num_fmt, n):
    """Return label string like '1.', 'a.', 'A.', 'i.' for a list item counter n (1-based)."""
    if num_fmt in (None, 'decimal'): return f'{n}.'
    if num_fmt == 'upperLetter': return chr(ord('A') + n - 1) + '.'
    if num_fmt == 'lowerLetter': return chr(ord('a') + n - 1) + '.'
    if num_fmt == 'upperRoman':  return _int_to_roman(n) + '.'
    if num_fmt == 'lowerRoman':  return _int_to_roman(n).lower() + '.'
    return '•'  # bullet / disc


def _make_hdr_prefix(num_fmt, idx):
    """Return prefix string like 'A. ' or 'I. ' for auto-numbered group headers."""
    n = idx + 1
    if num_fmt == 'upperLetter':
        return chr(ord('A') + idx) + '. '
    if num_fmt == 'lowerLetter':
        return chr(ord('a') + idx) + '. '
    if num_fmt == 'upperRoman':
        return _int_to_roman(n) + '. '
    if num_fmt == 'lowerRoman':
        return _int_to_roman(n).lower() + '. '
    return ''

def _get_num_fmt(doc, num_id, ilvl):
    """Return the w:numFmt val for a given numId+ilvl, or None if not found."""
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return None
        root = numbering_part._element
        # Locate <w:num w:numId="...">
        num_el = next(
            (n for n in root.findall(qn('w:num'))
             if n.get(qn('w:numId')) == str(num_id)),
            None
        )
        if num_el is None:
            return None
        abs_ref = num_el.find(qn('w:abstractNumId'))
        if abs_ref is None:
            return None
        abs_id = abs_ref.get(qn('w:val'))
        # Locate matching <w:abstractNum>
        abs_el = next(
            (a for a in root.findall(qn('w:abstractNum'))
             if a.get(qn('w:abstractNumId')) == abs_id),
            None
        )
        if abs_el is None:
            return None
        # Find the level entry
        for lvl in abs_el.findall(qn('w:lvl')):
            if lvl.get(qn('w:ilvl')) == str(ilvl):
                fmt_el = lvl.find(qn('w:numFmt'))
                if fmt_el is not None:
                    return fmt_el.get(qn('w:val'))
    except Exception:
        pass
    return None


# ── Section header keyword → section type mapping ────────────────────────────
_SECTION_KEYWORDS = {
    'TARGET LEARNING OUTCOMES': 'outcomes',
    'TARGET OUTCOMES':          'outcomes',
    'DATA AND RESULTS':         'data_sheet',
    'DATA SHEET':               'data_sheet',
    'DATASHEET':               'data_sheet',
    'DATASHEETS':               'data_sheet',
    'GUIDE QUESTIONS':          'guide_questions',
    'LABORATORY QUESTIONS':     'guide_questions',
    'INTRODUCTION':             'introduction',
    'BIBLIOGRAPHY':             'references',
    'REFERENCES':               'references',
    'OBJECTIVES':               'outcomes',
    'APPARATUS':                'materials',
    'EQUIPMENT':                'materials',
    'MATERIALS':                'materials',
    'MATERIAL':                 'materials',
    'PROCEDURE':                'procedures',
    'PROCEDURES':               'procedures',
    'QUESTIONS':                'guide_questions',
    'LABORATORY QUESTIONS AND PROBLEMS': 'guide_questions',
}


def is_section_header(text):
    """Return True if text matches any known section-header keyword."""
    upper = text.strip().upper()
    for kw in _SECTION_KEYWORDS:
        if kw in upper:
            return True
    return False


def _detect_section_type(upper_text):
    """Map a section header (uppercased) to its internal section-type string.
    Uses longest-match to handle overlaps (e.g. 'GUIDE QUESTIONS' > 'QUESTIONS').
    """
    upper = upper_text.strip()
    best_kw, best_type = '', 'introduction'
    for kw, stype in _SECTION_KEYWORDS.items():
        if kw in upper and len(kw) > len(best_kw):
            best_kw, best_type = kw, stype
    return best_type


def is_lab_title(text):
    return bool(re.match(r'LABORATORY\s+(ACTIVITY|EXERCISE|EXPERIMENT)', text.strip().upper()))


def is_bullet_point(para):
    style_name = para.style.name.lower() if para.style else ''
    if 'list' in style_name or 'bullet' in style_name:
        return True
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:numPr')) is not None:
        return True
    text = para.text.strip()
    if text and text[0] in ['•', '●', '◦', '▪', '▸', '\uf0b7']:
        return True
    return False


def is_numbered(text):
    # Matches: 1.  1)  1.1.  1.1)
    return bool(re.match(r'^\d+(?:\.\d+)*[\.\)]\s+', text))


def clean_bullet(text):
    return re.sub(r'^[\s\•\●\◦\▪\▸\uf0b7\-]+', '', text).strip()


def render_inline_equations(text):
    """
    Pass-through: equations are already formatted as LaTeX \\(...\\) / \\[...\\]
    by get_full_para_text, ready for MathJax rendering on the frontend.
    """
    return text


# ── Sub-section / group-header detection ─────────────────────────────────────
# Matches: "A.", "B.  Setting up", "B.1 Personal Error:", "1. Accuracy and Precision"
_SUBSEC_RE = re.compile(
    r'^('
    r'[A-Z]\.\s*$'                        # lone letter: "A."
    r'|[A-Z]\.\d+[\s\.]'                  # letter.digit: "B.1 ..." or "B.1."
    r'|[A-Z]\.\s+\S'                      # letter. word: "A. Setting up..."
    r'|\d+\.\s+[A-Z][a-zA-Z\s]{2,50}'    # numbered title: "1. Accuracy and Precision"
    r')',
    re.IGNORECASE
)

# ── Special content markers ───────────────────────────────────────────────────
_COMPUTATION_RE  = re.compile(r'show\s+(your\s+)?computations?', re.IGNORECASE)
_GRAPH_RE        = re.compile(r'attach\s+(graph|figure|plot|paper)\s+below', re.IGNORECASE)
_FIGURE_CAPTION_RE = re.compile(r'^(fig(ure)?\.?\s*\d+[\.\d]*[\.\:])', re.IGNORECASE)
# Bold prompts in data_sheet that look like answer prompts, not group headers
_DS_ANSWER_PROMPT_RE = re.compile(
    r'graph|attach|paste|picture|photo|image|figure|sketch|plot|draw|'
    r'show.{0,30}comput|computations?|calculations?|solve|derivat',
    re.IGNORECASE
)
_FILL_BLANK_RE   = re.compile(r'^_{2,}\s*(\d+[\.\)])\s*(.+)')
_SUB_QUESTION_RE = re.compile(r'^([a-z])[\.\)]\s+(.+)', re.DOTALL)
# Student info lines — skip these everywhere, already covered by built-in student info section
_STUDENT_INFO_RE = re.compile(
    r'^(names?(\s+or\s+group\s+members?)?|group\s+members?|course[/\s]|date\s+performed)',
    re.IGNORECASE
)


def _is_subsection_header(text):
    """Return True if text looks like a sub-section label inside intro/procedures."""
    return bool(_SUBSEC_RE.match(text.strip()))


def _parse_table_rows(table):
    """
    Parse a docx Table into structured rows.

    Each cell dict:
      {'type': 'text',  'value': '...'}  — cell has visible content
      {'type': 'input', 'value': ''}     — cell is blank → student fill-in

    Returns None if the table contains no visible content at all.
    """
    rows = []
    for row in table.rows:
        row_cells = []
        seen_tc_ids = set()
        for cell in row.cells:
            # python-docx returns the same cell object multiple times when cells are
            # merged across columns (colspan). Skip duplicates using the underlying
            # XML element identity so merged cells are counted only once.
            tc_id = id(cell._tc)
            if tc_id in seen_tc_ids:
                continue
            seen_tc_ids.add(tc_id)
            parts = []
            for para in cell.paragraphs:
                t = get_full_para_text(para)
                if t:
                    parts.append(render_inline_equations(t))
            cell_text = ' '.join(parts).strip()
            if not cell_text or re.match(r'^[_\-\s]*$', cell_text):
                row_cells.append({'type': 'input', 'value': ''})
            else:
                row_cells.append({'type': 'text', 'value': cell_text})
        rows.append(row_cells)
    if not any(c.get('value') for row in rows for c in row):
        return None
    return rows


def _rows_to_plain(rows):
    """Convert structured rows [{type,value}] to plain list-of-lists of strings."""
    return [[c.get('value', '') for c in row] for row in rows]


# ============================================================
# MAIN PARSE FUNCTION
# ============================================================

def parse_docx(filepath):
    doc = Document(filepath)
    images = extract_images(doc)
    body = doc.element.body

    # Map xml-element id → ('para', Paragraph) or ('table', Table).
    # doc.paragraphs / doc.tables return only direct body children (no nesting),
    # so this correctly captures document order without internal class imports.
    body_elem_map = {}
    for p in doc.paragraphs:
        body_elem_map[id(p._p)] = ('para', p)
    for t in doc.tables:
        body_elem_map[id(t._tbl)] = ('table', t)

    sections = []
    current_section = None
    current_content = []
    lab_title_done = False
    tables_data = []   # backward-compat flat list used by create_activity.html

    def save_section():
        if current_section is not None:
            content = current_content.copy()
            # Post-process guide_questions: nest tables + images under their preceding question
            if current_section['type'] == 'guide_questions':
                processed = []
                for item in content:
                    if item.get('type') in ('table', 'image'):
                        last_q = next((x for x in reversed(processed) if x.get('type') == 'question'), None)
                        if last_q is not None:
                            key = 'inline_tables' if item['type'] == 'table' else 'inline_images'
                            last_q.setdefault(key, []).append(item)
                        else:
                            processed.append(item)
                    else:
                        processed.append(item)
                content = processed
            # Post-process procedures: nest tables, images, and continuation text under their preceding step
            elif current_section['type'] == 'procedures':
                processed = []
                for item in content:
                    if item.get('type') in ('table', 'image'):
                        last_step = next((x for x in reversed(processed) if x.get('type') == 'step'), None)
                        if last_step is not None:
                            key = 'inline_tables' if item['type'] == 'table' else 'inline_images'
                            last_step.setdefault(key, []).append(item)
                        else:
                            processed.append(item)
                    elif item.get('type') in ('text', 'sub_step'):
                        text = item.get('text', '').strip()
                        last_step = next((x for x in reversed(processed) if x.get('type') == 'step'), None)
                        # Attach to preceding step if: equation line OR starts with lowercase (continuation)
                        is_eq_line = bool(re.match(r'^\s*\\\(', text)) or bool(re.match(r'^\s*\\\[', text))
                        is_continuation = bool(text) and (text[0].islower() or text[0] in ('(', '['))
                        if last_step is not None and (is_eq_line or is_continuation):
                            last_step.setdefault('continuation', []).append(item)
                        else:
                            processed.append(item)
                    else:
                        processed.append(item)
                content = processed
            sections.append({
                'type': current_section['type'],
                'title': current_section['title'],
                'content': content
            })

    for child in body:
        entry = body_elem_map.get(id(child))
        if entry is None:
            continue  # skip sectPr, bookmarks, etc.

        kind, obj = entry

        # ── TABLE (encountered inline within section) ─────────────────────
        if kind == 'table':
            if current_section is None:
                continue
            if current_section['type'] == 'lab_title' and lab_title_done:
                continue
            rows = _parse_table_rows(obj)
            if rows:
                tbl_item = {'type': 'table', 'rows': rows}
                # Detect table title: any short text (≤60 chars) immediately preceding the
                # table (looking past spacing items) is used as the title and removed from content.
                _NO_TITLE_TYPES = {
                    'table', 'list_item', 'outcome', 'step', 'question',
                    'image', 'image_group', 'figure_caption', 'spacing',
                    'computation_space', 'graph_attachment',
                }
                _look = len(current_content) - 1
                while _look >= 0 and current_content[_look].get('type') == 'spacing':
                    _look -= 1
                if _look >= 0:
                    _prev = current_content[_look]
                    if _prev.get('type', '') not in _NO_TITLE_TYPES:
                        _prev_text = _prev.get('text', '') or _prev.get('label', '')
                        if _prev_text and len(_prev_text) <= 100:
                            tbl_item['title'] = _prev_text
                            del current_content[_look:]   # remove title + any trailing spacings
                current_content.append(tbl_item)
            continue

        # ── PARAGRAPH ────────────────────────────────────────────────────
        para = obj
        raw_text = para.text.strip()
        has_omml = bool(para._element.find('.//' + qn('m:oMath')))

        if not raw_text and not get_para_image_ids(para) and not has_omml:
            # Preserve empty paragraphs as spacing inside Introduction
            if current_section and current_section['type'] == 'introduction':
                current_content.append({'type': 'spacing'})
            continue

        # Skip student info lines — redundant with built-in student info section
        if _STUDENT_INFO_RE.match(raw_text):
            continue

        bold = is_bold(para)

        # ── LAB TITLE ──
        if bold and is_lab_title(raw_text):
            save_section()
            current_section = {'type': 'lab_title', 'title': raw_text}
            current_content = []
            lab_title_done = False
            continue

        # ── LAB SUBTITLE (bold line right after lab_title header) ──
        if (current_section and current_section['type'] == 'lab_title'
                and not lab_title_done and bold and not is_section_header(raw_text)):
            current_content.append({'text': raw_text})
            continue

        # ── KNOWN SECTION HEADERS ──
        if bold and is_section_header(raw_text):
            if current_section and current_section['type'] == 'lab_title':
                lab_title_done = True
            save_section()
            stype = _detect_section_type(raw_text.upper())
            current_section = {'type': stype, 'title': raw_text}
            current_content = []
            continue

        # ── SECTION CONTENT ──
        if current_section is None:
            continue
        if current_section['type'] == 'lab_title' and lab_title_done:
            continue

        stype = current_section['type']

        # Images — group all images in same paragraph together
        img_ids = get_para_image_ids(para)
        if img_ids:
            srcs = [images[rid]['data'] for rid in img_ids if rid in images]
            if len(srcs) == 1:
                current_content.append({
                    'type': 'image',
                    'src': srcs[0],
                    'caption': raw_text or ''
                })
            elif len(srcs) > 1:
                current_content.append({
                    'type': 'image_group',
                    'images': [{'src': s} for s in srcs],
                    'caption': raw_text or ''
                })
            continue

        full_text = get_full_para_text(para)
        if not full_text:
            continue

        html_text = render_inline_equations(full_text)

        # ── INTRODUCTION ──
        if stype == 'introduction':
            pstyle = _para_style(para)
            numPr_el = para._element.find('.//' + qn('w:numPr'))
            # ── Auto-numbered / auto-bulleted list item ──
            if numPr_el is not None:
                ilvl_el  = numPr_el.find(qn('w:ilvl'))
                numId_el = numPr_el.find(qn('w:numId'))
                level  = int(ilvl_el.get(qn('w:val'), '0'))  if ilvl_el  is not None else 0
                num_id = numId_el.get(qn('w:val'), '0')      if numId_el is not None else '0'
                nfmt   = _get_num_fmt(doc, num_id, level)
                ltype  = 'ordered' if nfmt in ('decimal', 'upperRoman', 'lowerRoman',
                                                'upperLetter', 'lowerLetter') else 'bullet'
                # Track counter per (numId, level) to generate correct labels
                if '_list_ctrs' not in current_section:
                    current_section['_list_ctrs'] = {}
                key = (num_id, level)
                # Reset deeper levels when this level increments
                for k in list(current_section['_list_ctrs']):
                    if k[0] == num_id and k[1] > level:
                        current_section['_list_ctrs'][k] = 0
                current_section['_list_ctrs'][key] = current_section['_list_ctrs'].get(key, 0) + 1
                ctr = current_section['_list_ctrs'][key]
                label = _make_list_label(nfmt, ctr)
                clean_text = re.sub(r'\s*\[\d+\]\s*$', '', html_text).strip()
                if clean_text:
                    current_content.append({
                        'type': 'list_item', 'text': clean_text,
                        'label': label, 'level': level, 'list_type': ltype,
                    })
            elif _FIGURE_CAPTION_RE.match(raw_text):
                current_content.append({'type': 'figure_caption', 'text': render_inline_equations(raw_text)})
            elif bold and _is_subsection_header(raw_text):
                current_content.append({'type': 'sub_header', 'text': raw_text, **pstyle})
            elif _COMPUTATION_RE.search(raw_text):
                current_content.append({'type': 'computation_space', **pstyle})
            elif _GRAPH_RE.search(raw_text):
                current_content.append({'type': 'graph_attachment', **pstyle})
            else:
                clean_text = re.sub(r'\s*\[\d+\]\s*$', '', html_text).strip()
                if not clean_text:
                    pass
                elif clean_text.strip().startswith('\\[') and clean_text.strip().endswith('\\]'):
                    current_content.append({'type': 'display_equation', 'text': clean_text})
                else:
                    # Manually typed bullet/numbered list item
                    man_bullet = re.match(r'^[\•\●\◦\▪\▸\-]\s+(.+)', clean_text, re.DOTALL)
                    man_num    = re.match(r'^([A-Z]?\d+(?:\.\d+)*[\.\)]|[A-Z][\.\)]|\d+[\.\)])\s+(.+)', clean_text, re.DOTALL)
                    if man_bullet:
                        current_content.append({
                            'type': 'list_item', 'text': man_bullet.group(1),
                            'label': '•', 'level': 0, 'list_type': 'bullet',
                        })
                    elif man_num:
                        lbl = man_num.group(1)
                        # Indent sub-items like B.1, B.2 one level deeper
                        lvl = 1 if re.match(r'^[A-Z]\.\d', lbl) else 0
                        current_content.append({
                            'type': 'list_item', 'text': man_num.group(2),
                            'label': lbl, 'level': lvl, 'list_type': 'ordered',
                        })
                    else:
                        current_content.append({'type': 'paragraph', 'text': clean_text, **pstyle})

        # ── OUTCOMES ──
        elif stype == 'outcomes':
            clean = clean_bullet(full_text.strip())
            if not clean:
                continue
            txt = render_inline_equations(clean)
            # Check for manually typed number prefix (e.g. "1.1. text" or "1. text")
            m = re.match(r'^([\d]+(?:\.\d+)*[\.\)])\s+', full_text.strip())
            # Check for Word auto-list numbering
            has_numpr = para._element.find('.//' + qn('w:numPr')) is not None
            if m:
                num = m.group(1).rstrip('.')
                body = render_inline_equations(clean_bullet(full_text.strip()[m.end():]))
                if body:
                    current_content.append({'type': 'outcome', 'number': num, 'text': body})
            elif has_numpr:
                idx = sum(1 for c in current_content if c.get('type') == 'outcome')
                current_content.append({'type': 'outcome', 'number': f'1.{idx + 1}', 'text': txt})
            else:
                current_content.append({'type': 'outcome_intro', 'text': txt})

        # ── MATERIALS ──
        elif stype == 'materials':
            is_precaution_label = bool(re.match(r'^precaution', full_text.strip(), re.IGNORECASE))
            if is_precaution_label:
                current_content.append({
                    'type': 'note',
                    'text': render_inline_equations(full_text.strip()),
                    'note': render_inline_equations(full_text.strip()),
                })
            elif is_bullet_point(para):
                clean = clean_bullet(full_text)
                if clean:
                    last = current_content[-1] if current_content else None
                    after_precaution = (last and last.get('type') == 'note' and
                                        re.match(r'^precaution', last.get('text', ''), re.IGNORECASE))
                    if after_precaution or (last and last.get('type') == 'note'):
                        current_content.append({
                            'type': 'note',
                            'text': render_inline_equations(clean),
                            'note': render_inline_equations(clean),
                        })
                    else:
                        current_content.append({
                            'type': 'item',
                            'text': render_inline_equations(clean),
                            'item': render_inline_equations(clean),
                            'is_checkbox': True,
                        })
            else:
                clean = re.sub(r'^[\s\-]+', '', full_text).strip()
                if clean:
                    current_content.append({
                        'type': 'note',
                        'text': render_inline_equations(clean),
                        'note': render_inline_equations(clean),
                    })

        # ── PROCEDURES ──
        elif stype == 'procedures':
            # Letter-prefix headers (A., B., C.) — literal text, always group headers.
            if _PROC_LETTER_HDR_RE.match(raw_text) or (bold and _is_subsection_header(raw_text)):
                current_content.append({'type': 'sub_header', 'text': raw_text})
                continue

            if _COMPUTATION_RE.search(raw_text):
                current_content.append({'type': 'computation_space'})
            elif _GRAPH_RE.search(raw_text):
                current_content.append({'type': 'graph_attachment'})
            else:
                pPr = para._element.find(qn('w:pPr'))
                numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
                is_auto_numbered = numPr is not None

                match = re.match(r'^(\d+)[\.\)]\s+(.+)', full_text, re.DOTALL)
                if match:
                    current_content.append({
                        'type': 'step',
                        'step_number': int(match.group(1)),
                        'instruction': render_inline_equations(match.group(2).strip()),
                        'text': render_inline_equations(match.group(2).strip()),
                    })
                elif is_auto_numbered:
                    # Check numFmt — letter/roman formats = group header, decimal = step
                    num_id_el = numPr.find(qn('w:numId'))
                    ilvl_el   = numPr.find(qn('w:ilvl'))
                    num_id    = num_id_el.get(qn('w:val')) if num_id_el is not None else '0'
                    ilvl      = ilvl_el.get(qn('w:val'))   if ilvl_el   is not None else '0'
                    num_fmt   = _get_num_fmt(doc, num_id, ilvl)
                    if num_fmt in _HEADER_NUM_FMTS:
                        clean = full_text.strip()
                        if clean:
                            hdr_idx = sum(1 for c in current_content if c.get('type') == 'sub_header')
                            prefix  = _make_hdr_prefix(num_fmt, hdr_idx)
                            current_content.append({'type': 'sub_header', 'text': prefix + clean})
                    else:
                        _last_hdr = max((i for i, c in enumerate(current_content) if c.get('type') == 'sub_header'), default=-1)
                        step_num = sum(1 for c in current_content[_last_hdr + 1:] if c.get('type') == 'step') + 1
                        clean = full_text.strip()
                        if clean:
                            current_content.append({
                                'type': 'step',
                                'step_number': step_num,
                                'instruction': render_inline_equations(clean),
                                'text': render_inline_equations(clean),
                            })
                elif is_bullet_point(para):
                    clean = clean_bullet(full_text)
                    if clean:
                        current_content.append({
                            'type': 'sub_step',
                            'text': render_inline_equations(clean),
                        })
                else:
                    clean = full_text.strip()
                    if clean:
                        current_content.append({
                            'type': 'text',
                            'text': render_inline_equations(clean),
                        })

        # ── GUIDE QUESTIONS ──
        elif stype == 'guide_questions':
            # Numbered sub-item: optional leading underscores/spaces + "1) text"
            # Handles: "_1) text", "___1) text", "1) text", underline-blank + "1) text"
            num_sub_early = re.match(r'^[_\s]*(\d+\))\s+(.+)', full_text.strip(), re.DOTALL)
            if num_sub_early:
                current_content.append({
                    'type': 'sub_question',
                    'label': num_sub_early.group(1),
                    'text': render_inline_equations(num_sub_early.group(2).strip()),
                    'lines': 2,
                })
                continue

            # Get numPr info
            numPr = para._element.find('.//' + qn('w:numPr'))
            num_id = None
            ilvl = 0
            if numPr is not None:
                ilvl_el = numPr.find(qn('w:ilvl'))
                ilvl = int(ilvl_el.get(qn('w:val'), '0')) if ilvl_el is not None else 0
                numId_el = numPr.find(qn('w:numId'))
                num_id = numId_el.get(qn('w:val')) if numId_el is not None else None

            # Lazy-init: first numId seen in this section = main question list
            if '_gq_main_num_id' not in current_section:
                current_section['_gq_main_num_id'] = None

            # Manually typed "a. / a) ..." — always sub-question
            sub_match = _SUB_QUESTION_RE.match(full_text)
            if sub_match:
                current_content.append({
                    'type': 'sub_question',
                    'label': sub_match.group(1),
                    'text': render_inline_equations(sub_match.group(2).strip()),
                    'lines': 2,
                })
                continue

            if numPr is not None:
                if ilvl >= 1:
                    # Deeper indent = sub-question regardless of numId
                    is_sub = True
                elif current_section['_gq_main_num_id'] is None:
                    # First list encountered = main question list
                    current_section['_gq_main_num_id'] = num_id
                    is_sub = False
                else:
                    # Same numId = main question; different numId = sub-question
                    is_sub = (num_id != current_section['_gq_main_num_id'])

                clean = full_text.strip()
                if not clean:
                    continue
                if is_sub:
                    last_q_idx = max((i for i, c2 in enumerate(current_content) if c2.get('type') == 'question'), default=-1)
                    sub_idx = sum(1 for c in current_content[last_q_idx + 1:] if c.get('type') == 'sub_question')
                    num_fmt_sub = _get_num_fmt(doc, num_id, ilvl) if num_id else None
                    if num_fmt_sub == 'decimal':
                        label = f'{sub_idx + 1})'
                    else:
                        label = chr(ord('a') + sub_idx)
                    current_content.append({
                        'type': 'sub_question',
                        'label': label,
                        'text': render_inline_equations(clean),
                        'lines': 2,
                    })
                else:
                    q_num = sum(1 for c in current_content if c.get('type') == 'question') + 1
                    current_content.append({
                        'type': 'question',
                        'question_number': q_num,
                        'question': render_inline_equations(clean),
                        'text': render_inline_equations(clean),
                        'lines': 3,
                    })
                continue

            # Manually typed "1) ..." with closing paren = numbered sub-question
            # (leading spaces from underline-blanks are stripped; "1." period = main question)
            # Manually typed numbered question: "1. ..."
            match = re.match(r'^(\d+)\.\s+(.+)', full_text, re.DOTALL)
            if match:
                current_content.append({
                    'type': 'question',
                    'question_number': int(match.group(1)),
                    'question': render_inline_equations(match.group(2).strip()),
                    'text': render_inline_equations(match.group(2).strip()),
                    'lines': 3,
                })
            else:
                clean = full_text.strip()
                # Skip standalone answer-space lines (3+ underscores, nothing else)
                if re.match(r'^_{3,}$', clean):
                    continue
                if clean:
                    current_content.append({
                        'type': 'text',
                        'text': render_inline_equations(clean),
                    })

        # ── REFERENCES ──
        elif stype == 'references':
            clean = clean_bullet(full_text)
            if clean:
                current_content.append({'type': 'reference', 'text': render_inline_equations(clean)})

        # ── DATA SHEET ──
        elif stype == 'data_sheet':
            clean = full_text.strip()
            field_parts = _parse_ds_field_para(para)
            if field_parts:
                for p in field_parts:
                    if p['kind'] == 'label':
                        p['text'] = render_inline_equations(p['text'])
                current_content.append({'type': 'ds_field', 'parts': field_parts})
            elif clean:
                numPr = para._element.find('.//' + qn('w:numPr'))
                is_auto_num = numPr is not None
                is_header = False
                prefix = ''
                if is_auto_num:
                    num_id_el = numPr.find(qn('w:numId'))
                    ilvl_el   = numPr.find(qn('w:ilvl'))
                    num_id    = num_id_el.get(qn('w:val')) if num_id_el is not None else '0'
                    ilvl      = ilvl_el.get(qn('w:val'))   if ilvl_el   is not None else '0'
                    num_fmt   = _get_num_fmt(doc, num_id, ilvl)
                    if num_fmt in _HEADER_NUM_FMTS:
                        hdr_idx = sum(1 for c in current_content if c.get('type') == 'ds_header')
                        prefix  = _make_hdr_prefix(num_fmt, hdr_idx)
                        is_header = True
                    elif num_fmt == 'decimal':
                        is_header = True
                _is_answer_prompt = bool(_DS_ANSWER_PROMPT_RE.search(clean))
                if not is_header and bold and not _is_answer_prompt:
                    is_header = True
                if not is_header and not _is_answer_prompt and _DS_HDR_PREFIX_RE.match(clean):
                    is_header = True
                if is_header:
                    current_content.append({'type': 'ds_header', 'text': render_inline_equations(prefix + clean)})
                else:
                    # Non-header short texts are potential table titles (ds_text)
                    current_content.append({'type': 'ds_text', 'text': render_inline_equations(clean)})

        # ── FALLBACK ──
        else:
            current_content.append({'type': 'text', 'text': html_text})

    save_section()

    # Rebuild tables_data from data_sheet section only (backward compat for create_activity.html)
    ds = next((s for s in sections if s['type'] == 'data_sheet'), None)
    tables_data = [
        {'rows': _rows_to_plain(item['rows']), 'title': item.get('title', '')}
        for item in (ds['content'] if ds else [])
        if item.get('type') == 'table'
    ]

    # Debug: show detected sections and table counts
    for s in sections:
        tcount = sum(1 for c in s['content'] if isinstance(c, dict) and c.get('type') == 'table')
        icount = sum(len(c.get('inline_tables', [])) for c in s['content'] if isinstance(c, dict))
        if tcount or icount:
            print(f"[TABLES] section={s['type']} flat_tables={tcount} inline_tables={icount}")
    print(f"[TABLES] data.tables count={len(tables_data)}")


    # ===== STUDENT INFO (always at end) =====
    sections = [s for s in sections if s['type'] != 'student_info']
    sections.append({
        'type': 'student_info',
        'title': 'STUDENT INFORMATION',
        'content': [
            {'field': 'Name', 'type': 'member'},
            {'field': 'Course/Year/Section', 'type': 'dropdown'},
            {'field': 'Date', 'type': 'date'}
        ]
    })

    return {
        'sections': sections,
        'tables': tables_data   # backward-compat for create_activity.html
    }